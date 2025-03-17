from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, JsonResponse
from .models import PrayerRequest, Appointment, ServiceAttendance, Resource,Sermon, SermonSeries, SermonTag
from community.models import Announcement, VolunteerApplication, VolunteerOpportunity, Survey, SurveyResponse, Poll, PollOption, PollVote, Testimonial
from django.db import models  # Add this import at the top
from django.utils import timezone
from django.contrib import messages
from .forms import PrayerRequestForm
from django.db.models import Q
from django.core.mail import send_mail
from django.template.loader import render_to_string
from django.conf import settings
from django.core.mail import EmailMultiAlternatives
from django.utils.html import strip_tags  # Import strip_tags
from django.contrib.auth.views import PasswordResetView
from django.contrib.auth.models import User
from django.utils.translation import gettext as _
from django.urls import reverse_lazy
from .forms import AppointmentResponseForm
from django.utils.timezone import now
from datetime import datetime, timedelta
from collections import defaultdict

import csv
from django.http import HttpResponse
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Pt
from contact.models import Notification
from django.http import FileResponse, Http404
from django.core.paginator import Paginator
from django.db.models import Q

# sermons/views.py
from django.views.generic import ListView, DetailView, TemplateView

def generate_doc(request):
    # Fetch attendance data
    today = now().date()
    start_of_week = today - timedelta(days=today.weekday())
    service_days = ['Sunday', 'Tuesday', 'Friday']

    weekly_attendance = ServiceAttendance.objects.filter(
        date__gte=start_of_week,
        date__lte=start_of_week + timedelta(days=6),
        day__in=service_days
    ).order_by('date')

    # Create a new Document
    doc = Document()
    
    # Add title to the document
    doc.add_heading(f"Service Attendance from {start_of_week.strftime('%B %d, %Y')} to {(start_of_week + timedelta(days=6)).strftime('%B %d, %Y')}", 0)

    # Add content for each record
    for record in weekly_attendance:
        doc.add_heading(f"Date: {record.date} - Day: {record.day}", level=1)
        doc.add_paragraph(f"Total Attendance: {record.total_attendance}, Males: {record.males}, Females: {record.females}")
        doc.add_paragraph(f"Absentees: Informed: {record.absent_informed}, Not Informed: {record.absent_not_informed}")
        doc.add_paragraph(f"Special Absences: Students: {record.students_absent}, Traveled: {record.traveled}, Sick: {record.sick}")

    # Save the document to a BytesIO buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Return the response as a downloadable DOCX file
    response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="attendance_report.docx"'
    return response

def generate_pdf(request):
    # Fetch attendance data
    today = now().date()
    start_of_week = today - timedelta(days=today.weekday())
    service_days = ['Sunday', 'Tuesday', 'Friday']

    weekly_attendance = ServiceAttendance.objects.filter(
        date__gte=start_of_week,
        date__lte=start_of_week + timedelta(days=6),
        day__in=service_days
    ).order_by('date')

    # Create a file-like buffer to receive PDF data
    buffer = BytesIO()
    
    # Create a PDF object using the buffer
    p = canvas.Canvas(buffer, pagesize=letter)
    
    # Add title and content to the PDF
    p.setFont("Helvetica", 14)
    p.drawString(100, 750, f"Service Attendance from {start_of_week.strftime('%B %d, %Y')} to {(start_of_week + timedelta(days=6)).strftime('%B %d, %Y')}")
    y_position = 730

    for record in weekly_attendance:
        p.drawString(100, y_position, f"Date: {record.date} - Day: {record.day}")
        y_position -= 20
        p.drawString(100, y_position, f"Total Attendance: {record.total_attendance}, Males: {record.males}, Females: {record.females}")
        y_position -= 20
        p.drawString(100, y_position, f"Absentees: Informed: {record.absent_informed}, Not Informed: {record.absent_not_informed}")
        y_position -= 20
        p.drawString(100, y_position, f"Special Absences: Students: {record.students_absent}, Traveled: {record.traveled}, Sick: {record.sick}")
        y_position -= 40

    p.showPage()
    p.save()

    # Get PDF data from buffer
    pdf_data = buffer.getvalue()
    buffer.close()

    # Return the response as a downloadable PDF
    response = HttpResponse(pdf_data, content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="attendance_report.pdf"'
    return response


def export_attendance_csv(request):
    """
    Exports the weekly service attendance records as a CSV file.
    """
    today = now().date()
    start_of_week = today - timedelta(days=today.weekday())  # Get Monday of the current week
    end_of_week = start_of_week + timedelta(days=6)  # Get Sunday of the current week

    # Fetch attendance records for the current week
    weekly_attendance = ServiceAttendance.objects.filter(date__range=[start_of_week, end_of_week]).order_by('date')

    # Create the HTTP response with CSV content type
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="weekly_attendance.csv"'

    writer = csv.writer(response)
    
    # Write the CSV headers
    writer.writerow([
        "Date", "Day", "Total Attendance", "Males", "Females", "Adults", "Children",
        "Visitors", "Absent (Informed)", "Absent (Not Informed)", "Students Absent", "Traveled", "Sick", "Notes"
    ])

    # Write each attendance record as a row
    for record in weekly_attendance:
        writer.writerow([
            record.date, record.day, record.total_attendance, record.males, record.females,
            record.adults, record.children, record.visitors, record.absent_informed,
            record.absent_not_informed, record.students_absent, record.traveled, record.sick,
            record.notes if record.notes else "-"
        ])

    return response



# def generate_attendance_history_doc(request):
#     # Fetch historical attendance data
#     today = now().date()
#     start_of_week = today - timedelta(days=today.weekday())
#     service_days = ['Sunday', 'Tuesday', 'Friday']

#     # Fetch all attendance records ordered by date
#     attendance_records = ServiceAttendance.objects.filter(day__in=service_days).order_by('date')

#     # Create a new DOCX Document
#     doc = Document()
#     doc.add_heading("Attendance History", 0)

#     # Add content for each record
#     for record in attendance_records:
#         doc.add_heading(f"Date: {record.date} - Day: {record.day}", level=1)
#         doc.add_paragraph(f"Total Attendance: {record.total_attendance}, Males: {record.males}, Females: {record.females}")
#         doc.add_paragraph(f"Absentees: Informed: {record.absent_informed}, Not Informed: {record.absent_not_informed}")
#         doc.add_paragraph(f"Special Absences: Students: {record.students_absent}, Traveled: {record.traveled}, Sick: {record.sick}")
    
#     # Save the document to a BytesIO buffer
#     buffer = BytesIO()
#     doc.save(buffer)
#     buffer.seek(0)

#     # Return the response as a downloadable DOCX file
#     response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
#     response['Content-Disposition'] = 'attachment; filename="attendance_history_report.docx"'
#     return response


# def generate_attendance_history_pdf(request):
#     # Fetch historical attendance data
#     today = now().date()
#     start_of_week = today - timedelta(days=today.weekday())
#     service_days = ['Sunday', 'Tuesday', 'Friday']

#     # Fetch all attendance records ordered by date
#     attendance_records = ServiceAttendance.objects.filter(day__in=service_days).order_by('date')

#     # Create a file-like buffer to receive PDF data
#     buffer = BytesIO()
#     p = canvas.Canvas(buffer, pagesize=letter)

#     # Set title for the PDF
#     p.setFont("Helvetica", 14)
#     p.drawString(100, 750, f"Attendance History")
#     y_position = 730

#     for record in attendance_records:
#         p.drawString(100, y_position, f"Date: {record.date} - Day: {record.day}")
#         y_position -= 20
#         p.drawString(100, y_position, f"Total Attendance: {record.total_attendance}, Males: {record.males}, Females: {record.females}")
#         y_position -= 20
#         p.drawString(100, y_position, f"Absentees: Informed: {record.absent_informed}, Not Informed: {record.absent_not_informed}")
#         y_position -= 20
#         p.drawString(100, y_position, f"Special Absences: Students: {record.students_absent}, Traveled: {record.traveled}, Sick: {record.sick}")
#         y_position -= 40
    
#     p.showPage()
#     p.save()

#     # Get PDF data from buffer
#     pdf_data = buffer.getvalue()
#     buffer.close()

#     # Return the response as a downloadable PDF
#     response = HttpResponse(pdf_data, content_type='application/pdf')
#     response['Content-Disposition'] = 'attachment; filename="attendance_history_report.pdf"'
#     return response


# def export_attendance_history_csv(request):
#     """
#     Exports the historical service attendance records as a CSV file.
#     """
#     today = now().date()
#     service_days = ['Sunday', 'Tuesday', 'Friday']
    
#     # Fetch all attendance records ordered by date
#     attendance_records = ServiceAttendance.objects.filter(day__in=service_days).order_by('date')

#     # Create the HTTP response with CSV content type
#     response = HttpResponse(content_type='text/csv')
#     response['Content-Disposition'] = 'attachment; filename="attendance_history.csv"'

#     writer = csv.writer(response)
    
#     # Write the CSV headers
#     writer.writerow([
#         "Date", "Day", "Total Attendance", "Males", "Females", "Adults", "Children",
#         "Visitors", "Absent (Informed)", "Absent (Not Informed)", "Students Absent", "Traveled", "Sick", "Notes"
#     ])

#     # Write each attendance record as a row
#     for record in attendance_records:
#         writer.writerow([
#             record.date, record.day, record.total_attendance, record.males, record.females,
#             record.adults, record.children, record.visitors, record.absent_informed,
#             record.absent_not_informed, record.students_absent, record.traveled, record.sick,
#             record.notes if record.notes else "-"
#         ])

#     return response



# Helper function to group records by week and day
def group_attendance_by_week_and_day(attendance_records):
    grouped_data = defaultdict(lambda: defaultdict(list))  # Nested defaultdict
    for record in attendance_records:
        # Get the start of the week for each record (Monday of that week)
        week_start = record.date - timedelta(days=record.date.weekday())
        grouped_data[week_start][record.day].append(record)
    return grouped_data

def generate_attendance_history_doc(request):
    # Fetch historical attendance data
    today = now().date()
    service_days = ['Sunday', 'Tuesday', 'Friday']

    # Fetch all attendance records ordered by date
    attendance_records = ServiceAttendance.objects.filter(day__in=service_days).order_by('date')

    # Group records by week and day
    grouped_data = group_attendance_by_week_and_day(attendance_records)

    # Create a new DOCX Document
    doc = Document()
    doc.add_heading("Attendance History", 0)

    # Add content for each group (week and day)
    for week_start, week_data in grouped_data.items():
        doc.add_heading(f"Week starting: {week_start.strftime('%B %d, %Y')}", level=1)
        
        for day, records in week_data.items():
            doc.add_heading(f"Day: {day}", level=2)

            # Loop through each record for that day and display attendance data
            for record in records:
                doc.add_paragraph(f"Date: {record.date.strftime('%B %d, %Y')}", style='BodyText')

                # Add a bullet list for attendance and absentee details
                p = doc.add_paragraph(style='ListBullet')
                p.add_run(f"Total Attendance: {record.total_attendance}, Males: {record.males}, Females: {record.females}")

                p = doc.add_paragraph(style='ListBullet')
                p.add_run(f"Absentees: Informed: {record.absent_informed}, Not Informed: {record.absent_not_informed}")

                p = doc.add_paragraph(style='ListBullet')
                p.add_run(f"Special Absences: Students: {record.students_absent}, Traveled: {record.traveled}, Sick: {record.sick}")
                
                # Add a table for attendance breakdown
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'

                # Add table headers
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Adults'
                hdr_cells[1].text = 'Children'
                hdr_cells[2].text = 'Visitors'
                hdr_cells[3].text = 'Notes'

                # Add table row with data
                row_cells = table.add_row().cells
                row_cells[0].text = str(record.adults)
                row_cells[1].text = str(record.children)
                row_cells[2].text = str(record.visitors)
                row_cells[3].text = record.notes if record.notes else "No notes"

                doc.add_paragraph()  # Add spacing between records

    # Save the document to a BytesIO buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Return the response as a downloadable DOCX file
    response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="attendance_history_report.docx"'
    return response



def generate_attendance_history_pdf(request):
    # Fetch historical attendance data
    today = now().date()
    service_days = ['Sunday', 'Tuesday', 'Friday']

    # Fetch all attendance records ordered by date
    attendance_records = ServiceAttendance.objects.filter(day__in=service_days).order_by('date')

    # Group records by week and day
    grouped_data = group_attendance_by_week_and_day(attendance_records)

    # Create a file-like buffer to receive PDF data
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)

    # Set title for the PDF
    p.setFont("Helvetica", 14)
    p.drawString(100, 750, f"Attendance History")
    y_position = 730

    # Add content for each group (week and day)
    for week_start, week_data in grouped_data.items():
        p.drawString(100, y_position, f"Week starting: {week_start.strftime('%B %d, %Y')}")
        y_position -= 20

        for day, records in week_data.items():
            p.drawString(100, y_position, f"Day: {day}")
            y_position -= 20
            for record in records:
                p.drawString(100, y_position, f"Date: {record.date}")
                y_position -= 20
                p.drawString(100, y_position, f"Total Attendance: {record.total_attendance}, Males: {record.males}, Females: {record.females}")
                y_position -= 20
                p.drawString(100, y_position, f"Absentees: Informed: {record.absent_informed}, Not Informed: {record.absent_not_informed}")
                y_position -= 20
                p.drawString(100, y_position, f"Special Absences: Students: {record.students_absent}, Traveled: {record.traveled}, Sick: {record.sick}")
                y_position -= 40

    p.showPage()
    p.save()

    # Get PDF data from buffer
    pdf_data = buffer.getvalue()
    buffer.close()

    # Return the response as a downloadable PDF
    response = HttpResponse(pdf_data, content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="attendance_history_report.pdf"'
    return response

def export_attendance_history_csv(request):
    """
    Exports the historical service attendance records as a CSV file.
    """
    today = now().date()
    service_days = ['Sunday', 'Tuesday', 'Friday']
    
    # Fetch all attendance records ordered by date
    attendance_records = ServiceAttendance.objects.filter(day__in=service_days).order_by('date')

    # Group records by week and day
    grouped_data = group_attendance_by_week_and_day(attendance_records)

    # Create the HTTP response with CSV content type
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="attendance_history.csv"'

    writer = csv.writer(response)

    # Write the CSV headers
    writer.writerow([
        "Week Start", "Day", "Date", "Total Attendance", "Males", "Females", "Adults", "Children",
        "Visitors", "Absent (Informed)", "Absent (Not Informed)", "Students Absent", "Traveled", "Sick", "Notes"
    ])

    # Write each attendance record as a row
    for week_start, week_data in grouped_data.items():
        for day, records in week_data.items():
            for record in records:
                writer.writerow([
                    week_start.strftime('%B %d, %Y'), record.day, record.date, record.total_attendance,
                    record.males, record.females, record.adults, record.children, record.visitors,
                    record.absent_informed, record.absent_not_informed, record.students_absent, record.traveled,
                    record.sick, record.notes if record.notes else "-"
                ])

    return response


# def member_appointments(request):
#     # Get all appointments
#     appointments = Appointment.objects.all()
    
#     # List to store appointments with roles assigned
#     roles_data = []

#     # Loop through each appointment and check the roles
#     for appointment in appointments:
#         roles = []

#         # Check each role field and compare it with the logged-in user's member object
#         if appointment.morning_devotion_leader == request.user.member:
#             roles.append('Morning Devotion Leader')
#         if appointment.sunday_service_leader == request.user.member:
#             roles.append('Sunday Service Leader')
#         if appointment.song_leader == request.user.member:
#             roles.append('Song Leader')
#         if appointment.bible_study_leader_sunday == request.user.member:
#             roles.append('Sunday Bible Study Leader')
#         if appointment.preacher == request.user.member:
#             roles.append('Preacher')
#         if appointment.first_prayer_leader == request.user.member:
#             roles.append('First Prayer Leader')
#         if appointment.second_prayer_leader == request.user.member:
#             roles.append('Second Prayer Leader')
#         if appointment.third_prayer_leader == request.user.member:
#             roles.append('Third Prayer Leader')
#         if appointment.lord_supper_leader == request.user.member:
#             roles.append('Lord\'s Supper Leader')
#         if request.user.member in appointment.lord_supper_helpers.all():
#             roles.append('Lord\'s Supper Helper')
#         if appointment.announcer == request.user.member:
#             roles.append('Announcer')
#         if appointment.bible_study_leader_tuesday == request.user.member:
#             roles.append('Tuesday Bible Study Leader')
#         if appointment.tuesday_service_leader == request.user.member:
#             roles.append('Tuesday Service Leader')
#         if appointment.friday_prayer_leader == request.user.member:
#             roles.append('Friday Prayer Leader')

#         # If the user has any role in this appointment, add it to the roles_data
#         if roles:
#             roles_data.append({'appointment': appointment, 'roles': roles})

#     # Render the template with the filtered appointments and roles
#     return render(request, 'worship/member_appointments.html', {
#         'appointments': roles_data,
#     })



@login_required
def member_appointments(request):
    """Displays the user's pending and completed appointments"""
    # appointments = Appointment.objects.filter(member=request.user.member).order_by('-date')
    
    """Retrieve only active (non-deleted) appointments for the logged-in member."""
    appointments = Appointment.objects.filter(member=request.user.member, is_deleted=False).order_by('-date')
    
    upcoming_appointments = appointments.filter(date__gte=now().date())  # Future dates
    past_appointments = appointments.filter(date__lt=now().date())  # Past dates
    
    # past_appointments = appointments.filter(status='Denied')  # Past dates
    
    context = {
        'upcoming_appointments': upcoming_appointments,
        'past_appointments': past_appointments,
        'appointments': appointments,
    }
    
    return render(request, 'worship/member_appointments.html', context)



# @login_required
# def respond_appointment(request, appointment_id):
#     """Handles approval or rejection of an appointment"""
#     appointment = get_object_or_404(Appointment, id=appointment_id, member=request.user.member)

#     if request.method == 'POST':
#         form = AppointmentResponseForm(request.POST, instance=appointment)
#         if form.is_valid():
#             form.save()
#             messages.success(request, f'Your response for {appointment.role.name} on {appointment.date} has been recorded.')
#             return redirect('user_appointments')
#     else:
#         form = AppointmentResponseForm(instance=appointment)

#     return render(request, 'worship/respond_appointment.html', {'form': form, 'appointment': appointment})


@login_required
def respond_appointment(request, appointment_id):
    """Handles approval or rejection of an appointment, sends an email notification,
    and creates a persistent notification with appointment details."""
    appointment = get_object_or_404(Appointment, id=appointment_id, member=request.user.member)

    if request.method == 'POST':
        form = AppointmentResponseForm(request.POST, instance=appointment)
        if form.is_valid():
            appointment = form.save()
            
            # Create a notification with appointment details
            Notification.objects.create(
                title="Appointment Response Recorded",
                content=(
                    f"Dear {appointment.member.user.get_full_name()}, your response for {appointment.role.name} on "
                    f"{appointment.day.name} on {appointment.date.strftime('%B %d, %Y')} has been recorded. "
                    "Please check your appointment section in your dashboard for more information."
                ),
                member=appointment.member
            )
            
            # Attempt to send an email notification
            try:
                send_appointment_email(appointment)
            except Exception as e:
                messages.warning(request, "Your response was recorded, but the email notification failed.")

            messages.success(
                request, 
                f"Your response for {appointment.role.name} on {appointment.date.strftime('%B %d, %Y')} has been recorded."
            )
            return redirect('worship:member_appointments')
    else:
        form = AppointmentResponseForm(instance=appointment)

    return render(request, 'worship/respond_appointment.html', {'form': form, 'appointment': appointment})


def send_appointment_email(appointment):
    """Sends an email notification about the user's response."""
    subject = "Appointment Response Notification"
    recipient_email = "samytest777@gmail.com"  # Admin email (change to dynamic if needed)
    user_email = get_user_email(appointment.member)

    context = {
        'member_name': appointment.member.user.get_full_name(),
        'role': appointment.role.name,
        'date': appointment.date,
        'day': appointment.day.name,
        'status': appointment.status,
        'reason': appointment.reason if appointment.status == "Denied" else None,
    }

    html_content = render_to_string('worship/emails/appointment_response_email.html', context)

    email = EmailMultiAlternatives(subject, "", "noreply@church.com", [recipient_email, user_email])
    email.attach_alternative(html_content, "text/html")
    email.send()


def get_user_email(member):
    """Safely retrieves the user's email or returns a fallback."""
    if member and member.user and member.user.email:
        return member.user.email
    return "default@church.com"  # Use a default email if the user has no email
    

@login_required
def delete_appointment(request, appointment_id):
    """Marks the appointment as deleted, but keeps it in the database for admin records."""
    appointment = get_object_or_404(Appointment, id=appointment_id, member=request.user.member)

    # Prevent deletion if status is pending
    if appointment.status == Appointment.STATUS_PENDING:
        messages.error(request, "❌ Pending appointments cannot be deleted.")
        return redirect('worship:member_appointments')

    # Prevent deletion if appointment is accepted and still upcoming
    if appointment.status == Appointment.STATUS_ACCEPTED and appointment.date >= now().date():
        messages.error(request, "⚠️ You can only delete accepted appointments after the date has passed.")
        return redirect('worship:member_appointments')

    try:
        # Bypass model validation by using update()
        Appointment.objects.filter(id=appointment_id).update(is_deleted=True)
        
        # Create a notification for the deletion of the appointment
        Notification.objects.create(
            title="Appointment Deleted",
            content=(
                f"Dear {appointment.member.user.get_full_name()}, your appointment as "
                f"{appointment.role.name} on {appointment.day.name} on "
                f"{appointment.date.strftime('%B %d, %Y')} has been deleted."
            ),
            member=appointment.member
        )
        
        messages.success(request, "✅ Your appointment has been successfully deleted.")
    
    except Exception as e:
        messages.error(request, f"⚠️ Unexpected error: {e}")

    return redirect('worship:member_appointments')




@login_required
def prayer_requests(request):
    """View to list all prayer requests for the logged-in user."""
    # Filter prayer requests based on the logged-in user's member object
    prayer_requests = PrayerRequest.objects.filter(member=request.user.member).order_by('-date_requested')
    return render(request, 'worship/prayer_requests.html', {'prayer_requests': prayer_requests})



@login_required
def prayer_request_create(request):
    """Handles the creation of a new prayer request."""
    # Fetch only the logged-in user's prayer requests
    prayer_requests = PrayerRequest.objects.filter(member=request.user.member)  
    
    if request.method == 'POST':
        form = PrayerRequestForm(request.POST)
        if form.is_valid():
            prayer_request = form.save(commit=False)
            
            # Ensure the user has a related Member object
            if hasattr(request.user, 'member'):
                prayer_request.member = request.user.member  # Assign the current member
                prayer_request.save()
                
                # Create a notification for the submitted prayer request
                Notification.objects.create(
                    title="Prayer Request Submitted",
                    content=(
                        f"Dear {request.user.member.user.get_full_name()}, your prayer request has been submitted successfully. "
                        "Thank you for sharing your prayer needs with us."
                    ),
                    member=request.user.member
                )
                
                messages.success(request, 'Your prayer request has been submitted successfully!')
                return redirect('worship:prayer_requests')  # Redirect to the prayer request list
            else:
                messages.error(request, "You must be a registered member to submit a prayer request.")
                return redirect('worship:prayer_request_create')
        else:
            messages.error(request, "There was an error with the form. Please check and try again.")
    else:
        form = PrayerRequestForm()
    
    return render(request, 'worship/prayer_request_form.html', {
        'form': form,
        'prayer_requests': prayer_requests,  # Pass the queryset to the template
    })



@login_required
def prayer_request_update(request, pk):
    """Update an existing prayer request."""
    prayer_request = get_object_or_404(PrayerRequest, pk=pk, member=request.user.member)

    if request.method == 'POST':
        form = PrayerRequestForm(request.POST, instance=prayer_request)
        if form.is_valid():
            form.save()
            
            # Create a notification to inform the member about the update.
            Notification.objects.create(
                title="Prayer Request Updated",
                content=(
                    f"Dear {request.user.member.user.get_full_name()}, your prayer request has been updated successfully."
                ),
                member=request.user.member
            )
            
            messages.success(request, 'Prayer request updated successfully!')
            return redirect('worship:prayer_requests')
        else:
            messages.error(request, "There was an error updating your request, please check.")
    else:
        form = PrayerRequestForm(instance=prayer_request)

    return render(request, 'worship/prayer_request_form.html', {'form': form, 'editing': True})



@login_required
def prayer_request_delete(request, pk):
    prayer_request = get_object_or_404(PrayerRequest, pk=pk, member=request.user.member)

    if request.method == 'POST':
        prayer_request.delete()
        
        # Create a notification for the deletion of the prayer request
        Notification.objects.create(
            title="Prayer Request Deleted",
            content=(
                f"Dear {request.user.member.user.get_full_name()}, your prayer request has been deleted successfully."
            ),
            member=request.user.member
        )
        
        messages.success(request, 'Prayer request deleted successfully!')
        return redirect('worship:prayer_requests')

    return redirect('worship:prayer_requests')





def get_upcoming_week_start():
    today = datetime.today()
    days_until_sunday = (6 - today.weekday() + 1) % 7  # Days until next Sunday
    upcoming_sunday = today + timedelta(days=days_until_sunday)
    return upcoming_sunday.strftime("%B %d, %Y")  # Example: "February 18, 2025"

@login_required
def weekly_service_appointments(request):
    """
    Lists appointments grouped by service days (Sunday, Tuesday, Friday) for the current week.
    """
    today = now().date()
    start_of_week = today - timedelta(days=today.weekday())  # Get Monday of the current week

    # Define service days
    service_days = ['Sunday', 'Tuesday', 'Friday']
    
    # Fetch appointments for the current week
    weekly_appointments = Appointment.objects.filter(
        date__gte=start_of_week,
        date__lte=start_of_week + timedelta(days=6),
        day__name__in=service_days,
        is_archived=False  # Show only current week’s appointments
    ).order_by('date', 'role__name')

    # Group appointments by service day
    grouped_appointments = defaultdict(list)
    has_appointment = False  # ✅ Initialize to False
    
    for appointment in weekly_appointments:
        grouped_appointments[appointment.day.name].append(appointment)
        
      # If the logged-in user has an appointment, set flag to True
        if appointment.member == request.user.member:
            has_appointment = True

    context = {
        'grouped_appointments': dict(grouped_appointments),
        'upcoming_week_start': get_upcoming_week_start(),  # Fixed function here
        'has_appointment': has_appointment
    }
    return render(request, 'worship/weekly_appointments.html', context)



@login_required
def appointment_history(request):
    """
    Display past appointments grouped by week.
    """
    today = now().date()

    # Fetch past appointments (archived)
    past_appointments = Appointment.objects.filter(date__lt=today, is_archived=True).order_by('-date')

    # Group appointments by week
    grouped_appointments = defaultdict(list)
    for appointment in past_appointments:
        week_start = appointment.date - timedelta(days=appointment.date.weekday())  # Get Monday of the week
        week_label = week_start.strftime("%B %d, %Y")  # Example: "February 5, 2024"
        grouped_appointments[week_label].append(appointment)

    context = {
        'grouped_appointments': dict(grouped_appointments),
    }
    return render(request, 'worship/appointment_history.html', context)





# def appointment_history(request):
#     """
#     Display past appointments grouped by week and service days (Sunday, Tuesday, Friday).
#     """
#     today = now().date()
#     start_of_week = today - timedelta(days=today.weekday())  # Get Monday of the current week

#     # Define service days
#     service_days = ['Sunday', 'Tuesday', 'Friday']

#     # Fetch past appointments (archived) that match service days
#     past_appointments = Appointment.objects.filter(
#         date__lt=today,
#         is_archived=True,
#         day__name__in=service_days  # Ensure appointments fall on defined service days
#     ).order_by('-date')

#     # Group appointments by week and service day
#     grouped_appointments = defaultdict(lambda: defaultdict(list))
#     for appointment in past_appointments:
#         # Determine the start of the week for this appointment
#         appointment_week_start = appointment.date - timedelta(days=appointment.date.weekday())  
#         week_label = appointment_week_start.strftime("%B %d, %Y")  # Example: "February 5, 2024"

#         # Append appointment under the correct week and service day
#         grouped_appointments[week_label][appointment.day.name].append(appointment)

#     context = {
#         'grouped_appointments': dict(grouped_appointments),
#         'start_of_week': start_of_week.strftime("%B %d, %Y"),  # Display current week's Monday
#         'service_days': service_days  # Include service days in context
#     }
#     return render(request, 'worship/appointment_history.html', context)




# def weekly_service_attendance(request):
#     """
#     Lists service attendance records grouped by service days (Sunday, Tuesday, Friday) for the current week.
#     """
#     today = now().date()
#     start_of_week = today - timedelta(days=today.weekday())  # Get Monday of the current week

#     # Define service days
#     service_days = ['Sunday', 'Tuesday', 'Friday']
    
#     # Fetch attendance records for the current week
#     weekly_attendance = ServiceAttendance.objects.filter(
#         date__gte=start_of_week,
#         date__lte=start_of_week + timedelta(days=6),
#         day__in=service_days
#     ).order_by('date')

#     # Group attendance by service day
#     grouped_attendance = defaultdict(list)

#     for record in weekly_attendance:
#         grouped_attendance[record.day].append(record)

#     context = {
#         'grouped_attendance': dict(grouped_attendance),
#         'week_start': start_of_week.strftime('%B %d, %Y'),
#         'week_end': (start_of_week + timedelta(days=6)).strftime('%B %d, %Y')
#     }
#     return render(request, 'worship/weekly_attendance.html', context)

@login_required
def weekly_service_attendance(request):
    """
    Lists service attendance records grouped by service days (Sunday, Tuesday, Friday) for the current week.
    """
    today = now().date()
    start_of_week = today - timedelta(days=today.weekday())  # Get Monday of the current week
    end_of_week = start_of_week + timedelta(days=6)  # Sunday

    # Define service days
    service_days = ['Sunday', 'Tuesday', 'Friday']
    
    # Fetch attendance records for the current week
    weekly_attendance = ServiceAttendance.objects.filter(
        date__week=today.isocalendar()[1],  # Get records for the current week
        date__year=today.year,  # Ensure it's the current year
        day__in=service_days  # Match only service days
    ).order_by('date')

    # Group attendance by service day
    grouped_attendance = {day: [] for day in service_days}  # Ensure correct order
    for record in weekly_attendance:
        grouped_attendance[record.day].append(record)

    context = {
        'grouped_attendance': grouped_attendance,
        'week_start': start_of_week.strftime('%B %d, %Y'),
        'week_end': end_of_week.strftime('%B %d, %Y'),
    }
    return render(request, 'worship/weekly_attendance.html', context)


@login_required
def attendance_history(request):
    """
    Display past attendance records grouped by week, excluding the current week's attendance.
    """
    today = now().date()
    
    # Determine the start of the current week (Monday)
    week_start = today - timedelta(days=today.weekday())

    # Fetch attendance records before this week
    attendance_records = ServiceAttendance.objects.filter(date__lt=week_start).order_by('-date')

    # Group records by week
    grouped_attendance = defaultdict(list)
    for record in attendance_records:
        past_week_start = record.date - timedelta(days=record.date.weekday())  # Get Monday of that week
        week_label = past_week_start.strftime("%B %d, %Y")  # Example: "February 5, 2024"
        grouped_attendance[week_label].append(record)

    context = {
        'grouped_attendance': dict(grouped_attendance),  # Convert defaultdict to normal dict
    }

    return render(request, 'worship/attendance_history.html', context)



# @login_required
# def resource_list(request):
#     search_query = request.GET.get('q', '').strip()  # Trim whitespace
#     resources = Resource.objects.all()
    
#     if search_query:
#         resources = resources.filter(
#             Q(title__icontains=search_query) | 
#             Q(description__icontains=search_query)
#         )
    
#     resources = resources.order_by('-uploaded_at')
    
#     paginator = Paginator(resources, 10)
#     page_number = request.GET.get('page')
#     page_obj = paginator.get_page(page_number)
    
    
#     context = {
#         'resources': page_obj,
#         'search_query': search_query,
#     }
#       print(f"Debug: Context - {context}")  # Debugging
#     return render(request, 'worship/resource_list.html', context)




# @login_required
# def resource_list(request):
#     try:
#         print("Debug: View is being called!")  # Confirm the view is executed
#         search_query = request.GET.get('q', '').strip()
#         print(f"Debug: Search Query in View - '{search_query}'")

#         resources = Resource.objects.all()
#         if search_query:
#             resources = resources.filter(
#                 Q(title__icontains=search_query) | 
#                 Q(description__icontains=search_query)
#             )
        
#         resources = resources.order_by('-uploaded_at')
        
#         paginator = Paginator(resources, 10)
#         page_number = request.GET.get('page')
#         page_obj = paginator.get_page(page_number)
        
#         context = {
#             'resources': page_obj,
#             'resource_search_query': search_query,
#         }
#         print(f"Debug: Context - {context}")  # Debugging
#         return render(request, 'worship/resource_list.html', context)

#     except Exception as e:
#         print(f"Debug: Exception occurred - {e}")
#         raise  # Re-raise the exception for debugging



@login_required
def resource_list(request):
    try:
        print("Debug: View is being called!")  # Confirm the view is executed
        
        # Retrieve filter parameters from GET
        title_filter = request.GET.get('title', '').strip()
        description_filter = request.GET.get('description', '').strip()
        from_date = request.GET.get('from_date', '').strip()
        to_date = request.GET.get('to_date', '').strip()
        
        print(f"Debug: Filters - Title: '{title_filter}', Description: '{description_filter}', From: '{from_date}', To: '{to_date}'")
        
        resources = Resource.objects.all()
        
        # Apply filtering rules based on the provided parameters:
        if title_filter:
            resources = resources.filter(title__icontains=title_filter)
        if description_filter:
            resources = resources.filter(description__icontains=description_filter)
        if from_date:
            try:
                from_date_parsed = datetime.strptime(from_date, "%Y-%m-%d").date()
                resources = resources.filter(uploaded_at__date__gte=from_date_parsed)
            except Exception as e:
                print(f"Debug: Error parsing from_date - {e}")
        if to_date:
            try:
                to_date_parsed = datetime.strptime(to_date, "%Y-%m-%d").date()
                resources = resources.filter(uploaded_at__date__lte=to_date_parsed)
            except Exception as e:
                print(f"Debug: Error parsing to_date - {e}")
        
        resources = resources.order_by('-uploaded_at')
        
        # Pagination
        paginator = Paginator(resources, 2)
        page_number = request.GET.get('page')
        page_obj = paginator.get_page(page_number)
        
        context = {
            'resources': page_obj,
            'title_filter': title_filter,
            'description_filter': description_filter,
            'from_date': from_date,
            'to_date': to_date,
        }
        print(f"Debug: Context - {context}")  # Debugging output
        return render(request, 'worship/resource_list.html', context)

    except Exception as e:
        print(f"Debug: Exception occurred - {e}")
        raise  # Re-raise the exception for further debugging




@login_required
def resource_download(request, pk):
    """
    Allows a member to download a resource file as a PDF.
    """
    resource = get_object_or_404(Resource, pk=pk)
    if not resource.file:
        raise Http404("Resource file not found")
    
    # Open the file in binary mode and create a FileResponse
    response = FileResponse(resource.file.open('rb'), content_type='application/pdf')
    
    # Suggest a filename for the download
    response['Content-Disposition'] = f'attachment; filename="{resource.title}.pdf"'
    return response








# def appointment_history(request):
#     """
#     Display past appointments grouped by week.
#     """
#     today = now().date()

#     # Fetch past appointments (archived)
#     past_appointments = Appointment.objects.filter(date__lt=today, is_archived=True).order_by('-date')

#     # Group appointments by week
#     grouped_appointments = defaultdict(list)
#     for appointment in past_appointments:
#         week_start = appointment.date - timedelta(days=appointment.date.weekday())  # Get Monday of the week
#         week_label = week_start.strftime("%B %d, %Y")  # Example: "February 5, 2024"
#         grouped_appointments[week_label].append(appointment)

#     context = {
#         'grouped_appointments': dict(grouped_appointments),
#     }
#     return render(request, 'worship/appointment_history.html', context)

# List all sermons with featured, series, and tags (already used in your sermons archive)
# def sermons(request):
#     sermons_list = Sermon.objects.all().order_by('-date')
#     featured_sermon = Sermon.objects.filter(is_featured=True).order_by('-date').first()
#     latest_sermons = Sermon.objects.order_by('-date')[:6]  # Get the 6 latest sermons
#     series_list = SermonSeries.objects.all().order_by('-start_date')
#     tags_list = SermonTag.objects.all().order_by('name')
#     context = {
#         'sermons': sermons_list,
#         'featured_sermon': featured_sermon,
#         'latest_sermons': latest_sermons,
#         'series_list': series_list,
#         'tags_list': tags_list,
#     }
#     return render(request, 'worship/sermons.html', context)



def sermons(request):
    tag_slug = request.GET.get('tag')
    if tag_slug:
        selected_tag = get_object_or_404(SermonTag, slug=tag_slug)
        sermons_list = Sermon.objects.filter(tags=selected_tag).order_by('-date')
    else:
        selected_tag = None
        sermons_list = Sermon.objects.all().order_by('-date')

    featured_sermon = Sermon.objects.filter(is_featured=True).order_by('-date').first()
    latest_sermons = Sermon.objects.order_by('-date')[:6]  # Get the 6 latest sermons
    series_list = SermonSeries.objects.all().order_by('-start_date')
    tags_list = SermonTag.objects.all().order_by('name')

    context = {
        'sermons': sermons_list,
        'featured_sermon': featured_sermon,
        'latest_sermons': latest_sermons,
        'series_list': series_list,
        'tags_list': tags_list,
        'selected_tag': selected_tag,
    }
    return render(request, 'worship/sermons.html', context)

def sermon_detail(request, slug):
    sermon = get_object_or_404(Sermon, slug=slug)
    related_sermons = Sermon.objects.none()  # Empty queryset by default
    
    if sermon.series.exists():
        series = sermon.series.first()
        related_sermons = series.sermons.exclude(id=sermon.id).order_by('-date')[:4]
    
    context = {
        'sermon': sermon,
        'related_sermons': related_sermons,
    }
    return render(request, 'worship/sermon_detail.html', context)


# Detail view for a sermon series: list sermons in the series
def series_detail(request, slug):
    series = get_object_or_404(SermonSeries, slug=slug)
    sermons = series.sermons.all().order_by('-date')
    context = {
        'series': series,
        'sermons': sermons,
    }
    return render(request, 'worship/series_detail.html', context)

# Detail view for a sermon tag: list sermons with that tag
def tag_detail(request, slug):
    tag = get_object_or_404(SermonTag, slug=slug)
    sermons = tag.sermons.all().order_by('-date')
    context = {
        'tag': tag,
        'sermons': sermons,
    }
    return render(request, 'worship/tag_detail.html', context)